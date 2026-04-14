#!/usr/bin/env python3
"""
ueberschriften_aus_pdf.py
Extrahiert Überschriften aus den Stadtwächter-PDFs anhand von Schriftgröße und Fettdruck.

Kriterien für Überschriften:
  - Schriftgröße >= 18 pt                    → Haupttitel  (z. B. "Der Stadt-Wächter")
  - Schriftgröße >= 12.5 pt  UND  Fett       → Artikelüberschrift

Filter (werden herausgefiltert):
  - "Seite" + Zahl  (z. B. "Seite 1")
  - Zeilen die mit "Datei:" beginnen
  - Zeilen mit eckigen Klammern  (z. B. "[Fortsetzung …]")
  - Reiner Zeitungstitel "Der Stadt-Wächter" (Boilerplate)
  - Zeilen die nur "ANZEIGEN" oder "REKLAME" enthalten
  - Länge < 5 oder > 120 Zeichen

Performance: Multiprocessing mit bis zu 8 parallelen Prozessen.
Skipped-PDFs werden in ergebnisse/ueberschriften_v2_skipped.log vermerkt.

Ausgabe:
  ergebnisse/ueberschriften_v3.csv
  ergebnisse/ueberschriften_v3.docx
"""

import os
import re
import csv
import glob
from collections import defaultdict
from multiprocessing import Pool, cpu_count

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Pfade ──────────────────────────────────────────────────────────────────────
PDF_ROOT  = "/home/nghm/Dokumente/Zeitungsanalyse/daten_pdf"
OUT_DIR   = "/home/nghm/Dokumente/Zeitungsanalyse/ergebnisse"
CSV_PATH  = os.path.join(OUT_DIR, "ueberschriften_v3.csv")
DOCX_PATH = os.path.join(OUT_DIR, "ueberschriften_v3.docx")
LOG_PATH  = os.path.join(OUT_DIR, "ueberschriften_v3_skipped.log")

os.makedirs(OUT_DIR, exist_ok=True)

# ── Schwellenwerte ─────────────────────────────────────────────────────────────
SIZE_HAUPTTITEL   = 18.0   # Zeitungstitel
SIZE_UEBERSCHRIFT = 12.5   # Artikelüberschrift (muss zusätzlich bold sein)
SIZE_FALLBACK     = 11.0   # Fallback: Bold bei Normalgröße, aber kurze Zeile


# ── Hilfsfunktionen ────────────────────────────────────────────────────────────

def parse_filename(filename):
    """Extrahiert Jahrgang und Ausgabennummer aus dem Dateinamen."""
    name = os.path.splitext(os.path.basename(filename))[0]

    year_match = re.search(r'(192[0-9]|193[0-9])', name)
    year = int(year_match.group(1)) if year_match else None

    # Ausgabennummer am Anfang: "25_26", "64_65.", "01.", "OCR 3 ", "OCR 3." …
    num_match = re.match(r'[^\d]*(\d+)[_.]?(\d+)?', name)
    if num_match:
        issue = f"{num_match.group(1)}-{num_match.group(2)}" if num_match.group(2) else num_match.group(1)
    else:
        issue = "?"

    return year, issue


def chars_to_lines(chars):
    """Gruppiert pdfplumber-Zeichen nach y-Position zu Zeilen."""
    if not chars:
        return []

    chars = sorted(chars, key=lambda c: (round(c['top'], 1), c['x0']))

    lines, current = [], []
    prev_top = None

    for c in chars:
        top = round(c['top'], 1)
        if prev_top is None:
            prev_top = top
        if abs(top - prev_top) > 2:
            if current:
                lines.append(current)
            current = [c]
            prev_top = top
        else:
            current.append(c)

    if current:
        lines.append(current)

    result = []
    for lc in lines:
        text = ''.join(c['text'] for c in lc).strip()
        if not text:
            continue
        sizes    = [c['size'] for c in lc]
        fonts    = [c['fontname'] for c in lc]
        max_size = max(sizes)
        has_bold = any('Bold' in f or 'bold' in f for f in fonts)
        result.append({'text': text, 'max_size': max_size, 'has_bold': has_bold})
    return result


# ── Filtermuster (kompiliert für Performance) ──────────────────────────────────
_FILTER_PATTERNS = [
    re.compile(r'^\s*Seite\s+\d+\s*$', re.IGNORECASE),              # "Seite 1"
    re.compile(r'^\s*Datei\s*:', re.IGNORECASE),                      # "Datei: …"
    re.compile(r'[\[\]]'),                                             # eckige Klammern
    re.compile(r'^\s*Der\s+Stadt[-\s]?Wächter\s*$', re.IGNORECASE),  # Boilerplate-Titel
    re.compile(r'^\s*(ANZEIGEN|REKLAME)\b.*$', re.IGNORECASE),       # Anzeigen-Rubriken
    re.compile(r'Organ\s+für\s+freie\s+Meinung', re.IGNORECASE),     # Boilerplate-Untertitel
    re.compile(r'^\s*Garantierte\s+Auflage', re.IGNORECASE),          # Auflagenvermerk
]

MIN_LEN          = 5
MAX_LEN          = 120
MAX_LEN_FALLBACK = 80   # Fallback-Überschriften dürfen max. 80 Zeichen lang sein

# Muster für Seiten-Labels (auch in Großbuchstaben): "Seite 1", "SEITE 1", "S. 2" etc.
_SEITE_PATTERN = re.compile(r'^\s*(Seite|SEITE|S\.)\s*\d+\s*$', re.IGNORECASE)


def is_filtered(text):
    """Gibt True zurück wenn der Text herausgefiltert werden soll."""
    if len(text) < MIN_LEN or len(text) > MAX_LEN:
        return True
    return any(p.search(text) for p in _FILTER_PATTERNS)


# Wörter am Zeilenanfang die auf Fließtext/Metadaten hinweisen
_FLIESSTEXT_START = re.compile(
    r'^(Von|Für|An\s+den|An\s+die|Mit\s+|Durch\s+|Aus\s+der|Aus\s+dem|'
    r'Herausgegeben|Erscheint|Verlag|Schriftleiter|Nummer\s+\d|'
    r'Preis\s*:|Einzelpreis|Anzeigenpreis|Geschäftsstelle)\b',
    re.IGNORECASE
)

# Kleinwörter am Zeilenende → Fließtext-Abbruch (Artikel, Präpositionen, Konjunktionen)
_FLIESSTEXT_ENDE = re.compile(
    r'\b(und|oder|der|die|das|ein|eine|einen|einem|einer|des|dem|den|'
    r'in|im|an|am|auf|aus|bei|bis|für|mit|nach|seit|von|vor|zu|zur|zum|'
    r'wie|als|dass|ob|wenn|weil|aber|doch|auch|noch|schon|'
    r'ist|war|hat|wird|wurde|werden|haben|sein|sich|er|sie|es|wir|ihr)\s*$',
    re.IGNORECASE
)

def is_filtered_fallback(text):
    """
    Strenge Zusatzprüfung für den 11pt-Bold-Fallback (Normalgröße).
    Gibt True zurück wenn die Zeile als Fließtext eingestuft wird.

    Kriterien für eine gültige Fallback-Überschrift:
      - Kürzer als 80 Zeichen
      - Keine eckigen Klammern
      - Kein reines Seiten-Label
      - Beginnt mit Großbuchstabe
      - Endet nicht mit einem Kleinwort (kein Fließtext-Abbruch)
      - Enthält kein Komma oder Doppelpunkt (Fließtext-Indikator)
    """
    if len(text) > MAX_LEN_FALLBACK:
        return True
    if re.search(r'[\[\]]', text):
        return True
    if _SEITE_PATTERN.match(text):
        return True
    # Muss mit Großbuchstabe beginnen
    if not re.match(r'^[A-ZÄÖÜ„"»«]', text):
        return True
    # Typische Fließtext/Metadaten-Anfänge herausfiltern
    if _FLIESSTEXT_START.match(text):
        return True
    # Darf kein Komma oder Doppelpunkt enthalten
    if re.search(r'[,:]', text):
        return True
    # Mehrere Sätze (Punkt/Ausrufezeichen + Leerzeichen + Großbuchstabe) = Fließtext
    if re.search(r'[.!?]\s+[A-ZÄÖÜ]', text):
        return True
    # Darf nicht mit einem Kleinwort enden (Fließtext-Abbruch)
    if _FLIESSTEXT_ENDE.search(text):
        return True
    return False


def classify_line(line):
    """
    Gibt 'haupttitel', 'artikel' oder None zurück.

    Priorität:
      1. >= 18pt                     → haupttitel
      2. >= 12.5pt + Bold            → artikel
      3. ~11pt + Bold + kurz (<= 80) → artikel (Fallback für Ausgaben ohne Größenvarianz)
    """
    text = line['text']
    if not re.search(r'[A-Za-zÄÖÜäöüß]', text):
        return None
    if is_filtered(text):
        return None
    if line['max_size'] >= SIZE_HAUPTTITEL:
        return 'haupttitel'
    if line['max_size'] >= SIZE_UEBERSCHRIFT and line['has_bold']:
        return 'artikel'
    # Fallback: Bold bei Normalgröße, aber nur kurze Zeilen
    if line['has_bold'] and not is_filtered_fallback(text):
        return 'artikel'
    return None


# ── Worker-Funktion (wird in separatem Prozess ausgeführt) ────────────────────

def process_pdf(pdf_path):
    """
    Verarbeitet eine PDF-Datei und gibt ein Dict zurück:
    {year, issue, filename, headings: [{text, typ}], skipped: bool, skip_reason: str}
    """
    year, issue = parse_filename(pdf_path)
    headings = []
    seen = set()
    skipped = False
    skip_reason = ''

    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_chars = 0
            for page in pdf.pages:
                chars = page.chars
                if not chars:
                    continue
                total_chars += len(chars)
                for line in chars_to_lines(chars):
                    typ = classify_line(line)
                    if typ:
                        text = re.sub(r'\s+', ' ', line['text']).strip()
                        if text not in seen:
                            seen.add(text)
                            headings.append({'text': text, 'typ': typ})

            # PDF überspringen wenn kein verwertbarer Text vorhanden
            if total_chars == 0:
                skipped = True
                skip_reason = 'Keine Zeichen extrahierbar (möglicherweise reines Scan-Bild)'
    except Exception as e:
        skipped = True
        skip_reason = f'Fehler beim Öffnen/Lesen: {e}'

    return {
        'year':        year,
        'issue':       issue,
        'filename':    os.path.basename(pdf_path),
        'headings':    headings,
        'skipped':     skipped,
        'skip_reason': skip_reason,
    }


# ── Hauptprogramm ──────────────────────────────────────────────────────────────

def main():
    pdf_files = sorted(glob.glob(os.path.join(PDF_ROOT, "**", "*.pdf"), recursive=True))
    total = len(pdf_files)
    print(f"Gefundene PDF-Dateien: {total}")

    workers = min(8, cpu_count())
    print(f"Parallele Prozesse: {workers}\n")

    all_rows = []
    skipped_log = []
    done = 0

    with Pool(processes=workers) as pool:
        for result in pool.imap_unordered(process_pdf, pdf_files):
            done += 1

            if result['skipped']:
                print(f"[{done:3d}/{total}] ÜBERSPRUNGEN  {result['filename']}"
                      f"  → {result['skip_reason']}")
                skipped_log.append(result)
                continue

            n = len(result['headings'])
            print(f"[{done:3d}/{total}] Jahr={result['year']}  Ausg.={result['issue']:>6}  "
                  f"{n:3d} Überschriften  {result['filename']}")

            for h in result['headings']:
                all_rows.append({
                    'jahrgang':     result['year'],
                    'ausgabe':      result['issue'],
                    'typ':          h['typ'],
                    'ueberschrift': h['text'],
                    'dateiname':    result['filename'],
                })

    # Nach Jahrgang + Ausgabe sortieren für geordnete Ausgabe
    def sort_key(row):
        y = row['jahrgang'] or 9999
        m = re.match(r'(\d+)', str(row['ausgabe']))
        n = int(m.group(1)) if m else 9999
        return (y, n)

    all_rows.sort(key=sort_key)
    print(f"\nGesamt: {len(all_rows)} Überschriften  |  {len(skipped_log)} PDFs übersprungen\n")

    # ── CSV ────────────────────────────────────────────────────────────────────
    with open(CSV_PATH, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(
            f, fieldnames=['jahrgang', 'ausgabe', 'typ', 'ueberschrift', 'dateiname'])
        writer.writeheader()
        writer.writerows(all_rows)
    print(f"CSV gespeichert:           {CSV_PATH}")

    # ── Skipped-Log ───────────────────────────────────────────────────────────
    with open(LOG_PATH, 'w', encoding='utf-8') as f:
        f.write(f"Übersprungene PDFs ({len(skipped_log)})\n")
        f.write("=" * 60 + "\n")
        for r in skipped_log:
            f.write(f"{r['filename']}\n  Grund: {r['skip_reason']}\n\n")
    print(f"Skipped-Log gespeichert:   {LOG_PATH}")

    # ── Word-Dokument ─────────────────────────────────────────────────────────
    doc = Document()

    title_par = doc.add_heading('Überschriften – Der Stadt-Wächter (1929–1931)', level=0)
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Gruppieren: year → issue → [texts]
    grouped = defaultdict(lambda: defaultdict(list))
    for row in all_rows:
        grouped[row['jahrgang']][row['ausgabe']].append(row['ueberschrift'])

    for year in sorted(k for k in grouped if k is not None):
        yh = doc.add_heading(f'=== Jahrgang {year} ===', level=1)
        yh.runs[0].font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

        def issue_key(iss):
            m = re.match(r'(\d+)', str(iss))
            return int(m.group(1)) if m else 9999

        for issue in sorted(grouped[year].keys(), key=issue_key):
            doc.add_heading(f'Ausgabe {issue}:', level=2)
            for text in grouped[year][issue]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(text).font.size = Pt(11)

    doc.save(DOCX_PATH)
    print(f"Word-Dokument gespeichert: {DOCX_PATH}")

    # ── Erste 30 Überschriften ─────────────────────────────────────────────────
    print("\n" + "=" * 72)
    print("ERSTE 30 GEFUNDENE ÜBERSCHRIFTEN")
    print("=" * 72)
    for row in all_rows[:30]:
        print(f"  [{row['jahrgang']} / Ausg. {str(row['ausgabe']):>6}]"
              f"  [{row['typ']:10}]  {row['ueberschrift']}")
    print("=" * 72)
    print(f"\nFertig. {len(all_rows)} Überschriften in CSV und DOCX gespeichert.")


if __name__ == '__main__':
    main()
